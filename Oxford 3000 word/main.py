"""
This project involves extracting words from a table in a DOCX file. 
I discovered 3,000 English words from Oxford and realized I could use them. 
First, I converted the PDF file to a Word document using my converter app. 
Then, I gathered all the words and saved them in the "first.txt" file. After removing certain entries like (A1, B1, v., adj. etc.) and saved on "second.txt", 
Finnaly, I translated them and saved on "The_Oxford_3000.txt".Now I can add this file to my Word App Which I developed before.
"""

import docx
from googletrans import Translator

# read a docx file take the words of the table 

class Read_table_on_word:
    def __init__(self,file_path) -> None:
        self.file_path = file_path
        self.tables = self.read_word_tables()
        self.write_to_txt()
    
    # read tables   
    def read_word_tables(self):
        doc = docx.Document(self.file_path)
        tables = doc.tables
        table_data = []
        for table in tables:
            table_data.append([[cell.text for cell in row.cells] for row in table.rows])
        return table_data


               
    # then write on a txt file one by one
    def write_to_txt(self):
        with open("first.txt","w") as file:
            for i in range(11):
                for j in range(len(self.tables[i])):
                    new = self.tables[i][j]
                    for p in new:
                        file.write(f"{p}\n")

# Translate them and finnaly save on "The_Oxford_3000.txt"
class Translate:
    def __init__(self,old_word,word_translate) -> None:
        self.translate(old_word,word_translate)
          
    def translate(self,old_word,word_translate):
        with open("The_Oxford_3000.txt","a", encoding="utf-8") as file:
            try:
                translator = Translator()  
                result=translator.translate(word_translate,src="en",dest="tr").text
                file.write(f"{old_word}, {result},\n")
            except:
                pass

# check and delete all things except words than translate it 
class Check:
    def __init__(self) -> None:
        self.open()

    def open(self):
        with open("first.txt","r") as file:
            lines = file.readlines()
            for i in lines:
                self.write(self.check(i))
                Translate(i.replace("\n",""),self.check(i))

    def check(self, word):
        my_list = ["v.","n.","adj.","prep.","adv.",",","det.","/","conj.","A2","B1","B2","A1","\n","ad","B"]
        for i in my_list:
            word = word.replace(i, "")
        return word
    
    def write(self,word):
        with open("second.txt","a") as file:
            file.write(f"{word}\n")
    

if __name__ == "__main__":
    tables = Read_table_on_word('The_Oxford_3000.docx')
    check = Check()
  
